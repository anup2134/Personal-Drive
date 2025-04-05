from langchain_community.document_loaders import PyPDFLoader
from langchain_core.documents.base import Document 
from docx import Document as DOCX

def read_pdf(path:str)->list[Document]:
    loader = PyPDFLoader(path)
    docs = loader.load()

    return docs

def read_docx(path:str):
    doc = DOCX(path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"

    doc = Document(page_content=text)
    docs = [doc]
    return docs

def read_text_file(path:str)->list[Document]:
    with open(path, 'r') as f:
        text = f.read()
    doc = Document(page_content=text)
    docs = [doc]

    return docs

def parse_file(path:str,content_type:str)->list[Document]:
    TYPES = [
        "application/pdf",  # PDF files
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
        "text/plain",  # Plain text files
    ]

    if content_type == TYPES[0]:
        docs = read_pdf(path)
    elif content_type == TYPES[1]:
        docs = read_docx(path)
    elif content_type == TYPES[2]:
        docs = read_text_file(path)
    
    return docs
